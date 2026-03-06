from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

def add_table_from_md(doc, lines):
    rows_data = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows_data.append(cells)
    if len(rows_data) < 2:
        return
    header = rows_data[0]
    data = [r for r in rows_data[1:] if not all(set(c) <= set('-| ') for c in r)]
    
    table = doc.add_table(rows=1+len(data), cols=len(header), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            if ci < len(table.columns):
                cell = table.rows[ri+1].cells[ci]
                cell.text = val
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=1)
            for run in p.runs:
                run.font.name = 'Times New Roman'
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
        elif line.strip().startswith('|') and '|' in line:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_from_md(doc, table_lines)
            continue
        elif line.strip() == '---':
            pass
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            p.style.font.name = 'Times New Roman'
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip().strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif line.strip():
            text = line.strip()
            bold_parts = re.findall(r'\*\*(.*?)\*\*', text)
            plain_text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            plain_text = re.sub(r'`([^`]+)`', r'\1', plain_text)
            p = doc.add_paragraph(plain_text)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        i += 1
    
    doc.save(docx_path)
    print(f'Saved: {docx_path}')

convert_md_to_docx(
    r'd:\vladexecute\proj\Росток\Договор_оказания_услуг.md',
    r'd:\vladexecute\proj\Росток\Договор_оказания_услуг.docx'
)
