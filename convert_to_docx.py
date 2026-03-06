from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

def add_table_from_md(doc, lines):
    rows_data = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows_data.append(cells)
    if len(rows_data) < 2:
        return
    header = rows_data[0]
    data = [r for r in rows_data[1:] if not all(set(c) <= set('-| ') for c in r)]
    
    table = doc.add_table(rows=1+len(data), cols=len(header), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            if ci < len(table.columns):
                cell = table.rows[ri+1].cells[ci]
                cell.text = val
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.strip().startswith('|') and '|' in line:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_from_md(doc, table_lines)
            continue
        elif line.strip().startswith('```') and not line.strip() == '```':
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
                p.paragraph_format.left_indent = Inches(0.3)
        elif line.strip() == '---':
            doc.add_paragraph('─' * 60)
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip().strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        elif line.strip():
            text = line.strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`([^`]+)`', r'\1', text)
            doc.add_paragraph(text)
        
        i += 1
    
    doc.save(docx_path)
    print(f'Saved: {docx_path}')

convert_md_to_docx(r'd:\vladexecute\proj\Росток\ТЗ_Контент_Завод_v4.md', r'd:\vladexecute\proj\Росток\ТЗ_Контент_Завод_v4.docx')
